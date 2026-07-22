import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

def update_chapter_3_properly():
    doc_path = '/home/naantuam/Documents/Template for the Entire Project.docx'
    doc = docx.Document(doc_path)

    # Find position of 'CHAPTER FOUR'
    ch4_p = None
    for p in doc.paragraphs:
        if p.text.strip().upper() == 'CHAPTER FOUR':
            ch4_p = p
            break

    if not ch4_p:
        print("Error: Could not find CHAPTER FOUR heading")
        return

    def add_h1(text):
        p = ch4_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    def add_h2(text):
        p = ch4_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    def add_h3(text):
        p = ch4_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    def add_body(text):
        p = ch4_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)

    def add_code(text):
        p = ch4_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        p.paragraph_format.line_spacing = 1.0
        for r in p.runs:
            r.font.name = 'Courier New'
            r.font.size = Pt(9.0)
            r.font.color.rgb = RGBColor(30, 30, 30)

    def add_tbl(headers, rows_data):
        # Create table element
        t = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        for c_idx, h in enumerate(headers):
            cell = t.rows[0].cells[c_idx]
            cell.text = h
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(3)
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
                r.bold = True

        for r_idx, r_data in enumerate(rows_data):
            row_cells = t.rows[r_idx + 1].cells
            for c_idx, val in enumerate(r_data):
                row_cells[c_idx].text = str(val)
                row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = row_cells[c_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 3, 4, 5] else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9.5)

        # Move table before ch4_p
        ch4_p._p.addprevious(t._tbl)
        add_body("")

    # --- WRITE CHAPTER 3 ---
    add_h1("CHAPTER THREE")
    add_h2("SYSTEM ANALYSIS AND DESIGN (MATERIALS AND METHODS)")

    add_h3("3.1 Introduction")
    add_body("This chapter presents the system analysis, architectural design, materials, and operational methods used to build and evaluate the hybrid sensor monitoring framework. Modern mobile operating systems like Android rely on hardware peripherals such as camera lenses, microphones, location services, and motion sensors to provide rich functionality. Although permission controls govern direct access to cameras and microphones, zero-permission side-channel polling, covert background access, and silent data transmission present significant security risks. Existing defensive solutions often require superuser root access, introduce high computational overhead that degrades battery longevity, or depend entirely on continuous cloud connectivity. To address these limitations, this research introduces an unobtrusive hybrid monitoring framework that decouples lightweight mobile sensor auditing from real-time threat evaluation executed on an Edge Analysis Server.")
    add_body("The primary objective of this chapter is to explain how the system was analyzed, designed, implemented, and evaluated. Specifically, the system development and empirical telemetry benchmarking were carried out over a twelve-week period from May 2026 to July 2026. All experimental data collection and testing procedures were conducted within the Cybersecurity Science Laboratory at the Federal University of Technology Minna, utilizing a dedicated hardware testbed connected via a local 802.11ac Wi-Fi network and USB tethered interface.")

    add_h3("3.2 Research Methodology Framework")
    add_body("The development of the hybrid sensor monitoring framework follows the Design Science Research Methodology (DSRM), an established engineering framework tailored for constructing and evaluating cybersecurity software artifacts. DSRM structures the inquiry into six continuous, sequential phases, ensuring both technical rigor and empirical reproducibility.")

    add_h3("3.2.1 Phase 1: Problem Identification and Motivation")
    add_body("In this phase, empirical limitations in existing Android sensor security frameworks were catalogued. Standard mobile permission models fail to mitigate zero-permission motion sensor side-channels and covert background access. Current dynamic detection tools either require intrusive device rooting or suffer from an observability gap during network disconnections.")

    add_h3("3.2.2 Phase 2: Definition of Solution Objectives")
    add_body("The primary objective established during this phase was to engineer a hybrid monitoring framework capable of auditing hardware sensor access in real time without root permissions, maintaining offline operation, and suppressing false alerts for legitimate system services. This phase directly supports the fourth objective of this study: to evaluate the framework by mapping observed sensor behaviors to relevant techniques in the MITRE ATT&CK Mobile Matrix for structured threat classification using a custom rules engine, simulated attack testing, and live sensor monitoring across camera, microphone, location, and motion sensors.")

    add_h3("3.2.3 Phase 3: Design and Development")
    add_body("During this phase, the decoupled client-server architecture was designed and implemented. The artifact incorporates three primary components: the Mobile Telemetry Agent for non-intrusive sensor state collection, the Edge Analysis Server featuring a four-technique threat scoring engine, and the Forensic Web Dashboard for interactive security visualization.")

    add_h3("3.2.4 Phase 4: Demonstration")
    add_body("The prototype was deployed across dual operational modes: Host-Side Wireless ADB Bridge mode and Native On-Device Termux mode. The framework was demonstrated by streaming live telemetry events from an Android 10 smartphone to the Edge Analysis Server under both normal operating conditions and simulated attack scenarios.")

    add_h3("3.2.5 Phase 5: Evaluation")
    add_body("The implemented framework was subjected to systematic evaluation across four key metrics: detection accuracy, false positive rate, event processing latency, and mobile device battery consumption. Synthetic attack scripts simulating covert audio recording, background GPS tracking, and multi-sensor surveillance were executed to measure detection efficacy.")

    add_h3("3.2.6 Phase 6: Communication")
    add_body("The final phase involves documenting the research methodology, system design schemas, threat rules matrix, and empirical findings within this thesis document and publishing the underlying software artifacts for academic peer review.")

    add_h3("3.3 Experimental Setup and Technical Specifications")
    add_body("To guarantee exact experimental reproducibility, the hardware testbed and software environment were standardized. Table 3.1 outlines the hardware specifications of the target mobile device and the host Edge Analysis Server.")

    add_h3("Table 3.1: Hardware Testbed and Environment Specifications")
    add_tbl(
        ["Component", "Device / System Role", "Technical Specification"],
        [
            ["Target Mobile Handset", "Telemetry Source", "Infinix X683 (MediaTek Helio G70, 4 GB RAM, Android 10, API 29)"],
            ["Host Workstation", "Edge Analysis Server", "Intel Core i7-10700 CPU @ 2.90GHz, 16 GB DDR4 RAM, Ubuntu 22.04 LTS x86_64"],
            ["Local Network Interface", "Data Transport", "Dedicated 802.11ac Wi-Fi Router (5 GHz band) & USB 3.0 Tethered Link"]
        ]
    )

    add_body("Table 3.2 details the complete software stack, library dependencies, and platform tool versions used to build and execute the hybrid monitoring system.")

    add_h3("Table 3.2: Software Stack and System Tooling Versions")
    add_tbl(
        ["Software Layer", "Tool / Library Name", "Version", "Role / Purpose"],
        [
            ["Server Runtime", "Node.js", "v20.11.0", "Asynchronous JavaScript event loop for Edge Analysis Server"],
            ["Frontend Framework", "React / Vite", "v18.2.0 / v5.1.0", "User interface rendering and build tooling for Forensic Dashboard"],
            ["Local Database", "SQLite", "v3.37.2", "On-device offline telemetry buffering on Mobile Telemetry Agent"],
            ["Edge Database", "PostgreSQL (Neon)", "v15.4 (Driver v0.8.0)", "Centralized session and threat alert storage on Edge Server"],
            ["Network Protocol", "ws (WebSocket)", "v8.16.0", "Bidirectional real-time event streaming between agent and server"],
            ["Android Bridge", "Platform-Tools (ADB)", "v34.0.5", "Host-side Android diagnostic querying interface"]
        ]
    )

    add_body("Mobile telemetry is gathered directly from core Android system services via diagnostic interfaces. Specifically, the Mobile Telemetry Agent polls sensorservice for accelerometer and ambient light polling rates, audio for active microphone recording tracks, camera for active camera service sessions, and location for GPS and network location requests.")

    add_h3("3.4 System Requirements and Target Justifications")
    add_body("System analysis established clear functional and non-functional requirements to guide the engineering of the monitoring framework. Table 3.3 lists the numbered functional requirements.")

    add_h3("Table 3.3: System Functional Requirements")
    add_tbl(
        ["ID", "Functional Requirement Description", "Target Component"],
        [
            ["FR-01", "Audit camera, microphone, location, and motion sensors without requiring root superuser permissions.", "Mobile Telemetry Agent"],
            ["FR-02", "Package sensor events into structured JSON payloads containing app state and screen status.", "Mobile Telemetry Agent"],
            ["FR-03", "Evaluate incoming telemetry events in real time against a four-technique contextual threat matrix.", "Edge Analysis Server"],
            ["FR-04", "Apply a 5-minute cooldown throttle per application-sensor pair to prevent alert log flooding.", "Edge Analysis Server"],
            ["FR-05", "Buffer telemetry events in local SQLite storage during network connectivity disconnections.", "Mobile Telemetry Agent"],
            ["FR-06", "Render live threat streams, risk trends, and session details dynamically on an interactive UI.", "Forensic Web Dashboard"]
        ]
    )

    add_body("Table 3.4 outlines the non-functional requirements, including their measurable target thresholds and formal academic justifications.")

    add_h3("Table 3.4: Non-Functional Requirements and Target Justifications")
    add_tbl(
        ["ID", "Metric / Category", "Target Threshold", "Academic & Technical Justification"],
        [
            ["NFR-01", "Battery Overhead", "< 5% per hour", "Based on Android background service power budgets and prior research (Sikder et al., 2017) to prevent user-perceived battery drain."],
            ["NFR-02", "Event Latency", "< 100 ms", "Based on ITU-T network QoS guidelines for real-time security alerting to ensure immediate threat reporting."],
            ["NFR-03", "Offline Buffer", "100% Zero Loss", "Guarantees data integrity during network outages by caching up to 10,000 events locally before cloud sync."],
            ["NFR-04", "UI Responsiveness", "60 FPS Render", "Prevents DOM rendering lag during high-density alert streams using optimized React state management."],
            ["NFR-05", "CPU Utilization", "< 3% Host CPU", "Ensures the Edge Analysis Server maintains low resource utilization under continuous WebSocket polling."]
        ]
    )

    add_h3("3.5 System Architecture, Diagrams, and Operational Workflows")
    add_body("The hybrid monitoring framework uses a decoupled client-server architecture. The system comprises three main components: the Mobile Telemetry Agent, the Edge Analysis Server, and the Forensic Web Dashboard.")

    add_h3("3.5.1 System Architecture Topology")
    add_body("Figure 3.1 illustrates the structural topology of the hybrid monitoring system, showing data flow from Android system services to the edge server and forensic interface.")

    add_code(
"""+-----------------------------------------------------------------------------------+
|                               MOBILE HANDSET LAYER                                |
|  +--------------------+  +---------------------+  +----------------------------+  |
|  | Android Hardware   |  | Diagnostic Services |  | Mobile Telemetry Agent     |  |
|  | (Cam/Mic/GPS/Accel)|->| (dumpsys/app_ops)   |->| (Termux / Wireless ADB)    |  |
|  +--------------------+  +---------------------+  +-------------+--------------+  |
+-----------------------------------------------------------------|-----------------+
                                                                  | WebSocket (JSON)
                                                                  v
+-----------------------------------------------------------------------------------+
|                             EDGE ANALYSIS SERVER LAYER                            |
|  +--------------------+  +---------------------+  +----------------------------+  |
|  | WebSocket Broker   |->| Contextual Threat   |->| Dual Persistence Engine    |  |
|  | (Port 4444)        |  | Scoring Engine      |  | (SQLite / PostgreSQL)      |  |
|  +--------------------+  +---------------------+  +-------------+--------------+  |
+-----------------------------------------------------------------|-----------------+
                                                                  | HTTPS / WS
                                                                  v
+-----------------------------------------------------------------------------------+
|                           FORENSIC WEB DASHBOARD LAYER                            |
|  +--------------------+  +---------------------+  +----------------------------+  |
|  | Risk Trend Chart   |  | Live Threat Feed    |  | Threat Detail Inspector    |  |
|  | (Recharts)         |  | (Alert Drawer)      |  | (MITRE Mapping View)       |  |
|  +--------------------+  +---------------------+  +----------------------------+  |
+-----------------------------------------------------------------------------------+
Figure 3.1: Hybrid Sensor Monitoring Framework High-Level Architecture Topology"""
    )

    add_h3("3.5.2 System Use Case Analysis")
    add_body("Figure 3.2 illustrates the core use cases supported by the framework across the three system actors: Mobile Handset, Edge Analysis Server, and Security Analyst.")

    add_code(
"""                   +-----------------------------------------------+
                   |           HYBRID MONITORING SYSTEM            |
                   |                                               |
  ( Mobile )------>| -- (UC-1: Capture Sensor Telemetry)           |
  ( Handset)       | -- (UC-2: Buffer Events Offline in SQLite)   |
                   |                                               |
  ( Edge   )------>| -- (UC-3: Evaluate Contextual Threat Score)  |
  ( Server )       | -- (UC-4: Apply GMS OS Infrastructure Exempt)|
                   | -- (UC-5: Throttle Duplicate Threat Alerts)   |
                   |                                               |
  ( Security)----->| -- (UC-6: View Live Threat Alerts on UI)      |
  ( Analyst )      | -- (UC-7: Inspect MITRE ATT&CK Mapping Detail)|
                   +-----------------------------------------------+
Figure 3.2: System Use Case Diagram for Hybrid Sensor Monitoring Framework"""
    )

    add_h3("3.5.3 Data Flow Diagram (DFD Level 1)")
    add_body("Figure 3.3 presents the Level 1 Data Flow Diagram, depicting how telemetry data travels through processing functions, rule evaluation, database storage, and dashboard display.")

    add_code(
"""[Android System] --(Raw dumpsys)--> (1.0 Telemetry Parsing) --(Structured JSON)--> [WebSocket Broker]
                                                                                            |
                                                                                    (Telemetry Packet)
                                                                                            v
[PostgreSQL DB] <--(Persist Log)-- (3.0 Rule Scoring Engine) <--(Context Event)-- (2.0 Context Extraction)
        |                                   |
(Stored Alerts)                     (Threat Alert)
        v                                   v
[Forensic UI]   <--(Render Feed)-- (4.0 Alert Throttling)
Figure 3.3: Level 1 Data Flow Diagram (DFD) for Telemetry Processing"""
    )

    add_h3("3.5.4 Telemetry Event Sequence Diagram")
    add_body("Figure 3.4 illustrates the message sequence for a telemetry event, from sensor query to threat scoring and dashboard update.")

    add_code(
"""Handset Agent            WebSocket Broker           Threat Engine           PostgreSQL DB            Forensic UI
      |                          |                        |                       |                       |
      |--1. Query dumpsys------->|                        |                       |                       |
      |--2. Send JSON Event----->|                        |                       |                       |
      |                          |--3. Evaluate Packet--->|                       |                       |
      |                          |                        |--4. Score Event------>|                       |
      |                          |                        |--5. Save Threat Log-->|                       |
      |                          |<--6. Threat Alert------|                       |                       |
      |                          |------------------------7. Broadcast Alert Stream-------------->|
      |                          |                        |                       |                       |
Figure 3.4: Sequence Diagram for Telemetry Processing and Threat Alerting"""
    )

    add_h3("3.5.5 Operational Deployment Modes")
    add_body("To support diverse operational environments, the Mobile Telemetry Agent provides two deployment modes: Host-Side Wireless ADB Bridge mode and Native On-Device Termux mode. In Wireless ADB Bridge mode, the Edge Analysis Server executes ADB shell commands over Wi-Fi to query handset state without installing custom applications on the device. In Native Termux mode, the agent runs directly within the device's shell environment.")

    add_h3("3.6 Technical Threat Engine, Algorithm, and MITRE ATT&CK Mapping")
    add_body("The Edge Analysis Server incorporates a four-technique threat evaluation engine that computes a composite risk score for every telemetry event.")

    add_h3("3.6.1 Four-Technique Threat Scoring Engine")
    add_body("The scoring model uses four analytical techniques, with point values derived from Android threat literature (Sikder et al., 2021; Muhammad et al., 2023) and refined through empirical testing:")
    add_body("Technique 1 (App Trust Tiers): Modifies risk based on application trust. OS infrastructure packages (such as Google Play Services) receive a -50 point reduction because background operation is architecturally expected. Trusted system applications receive -30 points, while known popular apps receive -15 points.")
    add_body("Technique 2 (Sensor-App Coherence): Evaluates whether sensor usage matches an application's declared purpose. Coherent usage (e.g. navigation app accessing GPS) receives a -10 point match reward, while incoherent usage (e.g. finance app accessing microphone) receives a +15 point mismatch penalty.")
    add_body("Technique 3 (Multi-Sensor Correlation): Detects compound surveillance patterns. Simultaneous camera and microphone usage triggers an Audio-Video sync penalty of +35 points. Simultaneous camera, microphone, and location access triggers a full surveillance triad penalty of +50 points.")
    add_body("Technique 4 (Origin Risk): Penalizes applications installed from non-Play Store sources (+20 points) due to elevated supply chain risk.")
    add_body("To prevent false alerts for system location brokers like Google Play Services (GMS), the engine enforces a 5-property OS Infrastructure Exemption. System brokers score 0 (Benign) because they possess platform signing certificates, act strictly as API brokers, undergo Play Protect verification, are tracked by system app-ops, and lack data exfiltration paths.")

    add_h3("3.6.2 Threat Evaluation Algorithm")
    add_body("Algorithm 3.1 presents the step-by-step logic executed by the Edge Analysis Server for incoming telemetry packets.")

    add_code(
"""Algorithm 3.1: Contextual Threat Evaluation Algorithm
Input  : Telemetry Packet P = { payload, metadata }
Output : Evaluation Result R = { TotalScore, ThreatLevel, TriggeredRules }

1. Initialize TotalScore = 0, TriggeredRules = []
2. Extract appPackage, sensorName, appState, screenState from P
3. Evaluate Base Context:
   If screenState == "OFF" Then TotalScore += 30; Append STATE_SCREEN_OFF
   ElseIf appState == "BACKGROUND" Then TotalScore += 25; Append STATE_BACKGROUND
4. Evaluate Sensor Activity (Mic (+20), Camera (+20), GPS (+15), BLE (+15))
5. If High-Value Sensor Active AND (appState == "BACKGROUND" OR screenState == "OFF"):
   TotalScore += 50; Append CONTEXT_BG_VIOLATION
6. Evaluate Trust Tier:
   If appPackage in OS_INFRASTRUCTURE:
      TotalScore -= 50; Append TRUST_OS_INFRA
      Remove CONTEXT_BG_VIOLATION penalty (OS Infra Exemption)
   ElseIf appPackage in TRUSTED_SYSTEM: TotalScore -= 30
   ElseIf appPackage in KNOWN_APP: TotalScore -= 15
7. Evaluate Coherence & Multi-Sensor Correlation (AV Sync (+35), Trio (+50))
8. Evaluate Origin Risk: If Sideloaded Then TotalScore += 20; Append ORIGIN_SIDELOADED
9. Compute Final Threat Level:
   If TotalScore >= 100 Then ThreatLevel = "CRITICAL"
   ElseIf TotalScore >= 60 Then ThreatLevel = "HIGH"
   ElseIf TotalScore >= 25 Then ThreatLevel = "SUSPICIOUS"
   Else ThreatLevel = "BENIGN"
10. Return { TotalScore, ThreatLevel, TriggeredRules }"""
    )

    add_h3("3.6.3 MITRE ATT&CK Mobile Threat Matrix")
    add_body("Table 3.5 maps the framework's threat rules to corresponding MITRE ATT&CK Mobile techniques.")

    add_h3("Table 3.5: Threat Matrix Mapped to MITRE ATT&CK Mobile Framework")
    add_tbl(
        ["Tactical Domain", "Rule ID", "Sensor Behavior & Trigger Condition", "MITRE ID", "MITRE Technique Name", "Points"],
        [
            ["Collection", "COLLECTION_MIC", "Active microphone recording stream", "T1430", "Audio Capture", "+20"],
            ["Collection", "COLLECTION_CAMERA", "Active camera capture session", "T1125", "Video Capture", "+20"],
            ["Collection", "COLLECTION_LOW_FREQ", "Motion sensor polling (<= 20 Hz)", "T1636", "Sensor Information Discovery", "+5"],
            ["Collection", "COLLECTION_HIGH_FREQ", "High-frequency motion polling (> 100 Hz)", "T1429", "Side-Channel Exploitation", "+20"],
            ["Discovery", "DISCOVERY_GPS", "GPS location provider polling", "T1636", "Location Tracking", "+15"],
            ["Discovery", "DISCOVERY_BLE", "Bluetooth Low Energy background scan", "T1636", "Sensor Information Discovery", "+15"],
            ["Discovery", "DISCOVERY_FUSION", "Simultaneous motion sensor reads", "T1427", "System Information Discovery", "+20"],
            ["Context", "STATE_BACKGROUND", "Sensor access in background state", "N/A", "Contextual State Penalty", "+25"],
            ["Context", "STATE_SCREEN_OFF", "Sensor streaming while screen is OFF", "N/A", "Covert Execution Context", "+30"],
            ["Context", "CONTEXT_BG_VIOLATION", "High-value sensor accessed in background", "N/A", "Unsanctioned Access", "+50"],
            ["Defense Evasion", "EVASION_STATE_DISCREPANCY", "Foreground claim mismatches display OFF", "T1036", "Masquerading / State Mismatch", "+20"],
            ["Defense Evasion", "EVASION_ACCESSIBILITY", "Non-system app using Accessibility", "T1406", "Abuse Accessibility Features", "+15"],
            ["Defense Evasion", "EVASION_LOG_DELETE", "Local log file deletion attempted", "T1403", "Indicator Removal on Host", "+30"],
            ["Exfiltration", "EXFIL_STAGING", "Large local JSON buffer staging detected", "T1430", "Data Staged", "+15"],
            ["Exfiltration", "EXFIL_IMMEDIATE", "Immediate exfiltration over WebSocket C2", "T1041", "Exfiltration Over C2 Channel", "+40"],
            ["Correlation", "CORRELATION_AV_SYNC", "Camera and Microphone active together", "T1512", "Audio & Video Capture Sync", "+35"],
            ["Correlation", "CORRELATION_TRACK_TRIO", "Camera + Mic + Location active together", "T1430", "Full Surveillance Profile", "+50"],
            ["Correlation", "CORRELATION_NET_SENSOR", "Sensor collection with active exfiltration", "T1041", "Exfiltrated Sensor Capture", "+25"],
            ["Origin Risk", "ORIGIN_SIDELOADED", "Package installed from non-Play Store origin", "T1476", "Sideloaded Application", "+20"]
        ]
    )

    add_h3("3.7 Database Design, Schema, and Synchronization Protocols")
    add_body("The framework employs a dual-database storage design combining local on-device buffering with centralized edge storage.")

    add_h3("3.7.1 Database Entity-Relationship (ER) Architecture")
    add_body("Figure 3.5 illustrates the relational entity structure between sessions, telemetry events, and threat alerts.")

    add_code(
"""+-----------------------+       1:N       +-----------------------+       1:1       +-----------------------+
|       SESSIONS        |---------------->|   TELEMETRY_EVENTS    |---------------->|     THREAT_ALERTS     |
+-----------------------+                 +-----------------------+                 +-----------------------+
| PK: id                |                 | PK: id                |                 | PK: id                |
| device_id: VARCHAR    |                 | FK: session_id        |                 | FK: telemetry_event_id|
| connection_type:VARCHAR                 | app_package: VARCHAR  |                 | threat_level: VARCHAR |
| created_at: TIMESTAMP |                 | sensor_name: VARCHAR  |                 | total_score: INTEGER  |
+-----------------------+                 | app_state: VARCHAR    |                 | triggered_rules: TEXT |
                                          | timestamp: TIMESTAMP  |                 | created_at: TIMESTAMP |
                                          +-----------------------+                 +-----------------------+
Figure 3.5: Entity-Relationship (ER) Schema Diagram"""
    )

    add_h3("3.7.2 Database Schema Definitions")
    add_body("Table 3.6 details the relational database schema definitions implemented across SQLite and PostgreSQL.")

    add_h3("Table 3.6: Relational Database Schema Specifications")
    add_tbl(
        ["Table Name", "Column Name", "Data Type", "Constraints", "Description"],
        [
            ["sessions", "id", "INTEGER / SERIAL", "PRIMARY KEY", "Unique session identifier"],
            ["sessions", "device_id", "VARCHAR(100)", "NOT NULL", "Hardware device model identifier"],
            ["sessions", "connection_type", "VARCHAR(50)", "NOT NULL", "Connection mode (wireless_adb / termux)"],
            ["telemetry_events", "id", "INTEGER / SERIAL", "PRIMARY KEY", "Unique event record identifier"],
            ["telemetry_events", "session_id", "INTEGER", "FOREIGN KEY", "References sessions(id)"],
            ["telemetry_events", "app_package", "VARCHAR(150)", "NOT NULL", "Android application package name"],
            ["telemetry_events", "sensor_name", "VARCHAR(50)", "NOT NULL", "Target hardware sensor queried"],
            ["threat_alerts", "id", "INTEGER / SERIAL", "PRIMARY KEY", "Unique threat alert identifier"],
            ["threat_alerts", "threat_level", "VARCHAR(20)", "NOT NULL", "Classified threat level (CRITICAL/HIGH)"],
            ["threat_alerts", "total_score", "INTEGER", "NOT NULL", "Computed composite threat risk score"]
        ]
    )

    add_h3("3.7.3 Asynchronous Offline Buffer Synchronization Routine")
    add_body("If network connectivity to the Edge Analysis Server is lost, the Mobile Telemetry Agent buffers telemetry events locally in SQLite. Upon network restoration, an automated synchronization routine uploads buffered records to Neon PostgreSQL in batch transactions and purges the local queue, guaranteeing zero data loss.")

    add_h3("3.8 System Evaluation Plan")
    add_body("To validate the framework's effectiveness, a comprehensive system evaluation plan was designed based on four performance criteria.")

    add_h3("3.8.1 Evaluation Metrics")
    add_body("1. Detection Accuracy and False Positive Rate (FPR): Evaluated by executing synthetic attack scripts and calculating True Positive (TP), False Positive (FP), True Negative (TN), and False Negative (FN) classifications.")
    add_body("2. Event Processing Latency: Measured as the round-trip time (in milliseconds) from initial dumpsys query on the handset to threat alert rendering on the dashboard.")
    add_body("3. Mobile Battery Overhead: Measured across 1-hour and 6-hour continuous polling intervals using Android Battery Historian to confirm compliance with NFR-01 (< 5% per hour).")

    add_h3("3.8.2 Synthetic Attack Test Suite")
    add_body("Table 3.7 lists the synthetic attack scenarios designed to evaluate the threat engine.")

    add_h3("Table 3.7: Synthetic Attack Test Suite for System Evaluation")
    add_tbl(
        ["Test ID", "Attack Scenario Name", "Simulated Threat Behavior", "Expected Threat Classification"],
        [
            ["TC-01", "Covert Audio Recording", "Background microphone access with display screen OFF", "CRITICAL (Score >= 100)"],
            ["TC-02", "Background GPS Tracking", "Unsanctioned location access by third-party app in background", "HIGH (Score >= 60)"],
            ["TC-03", "Surveillance Triad", "Simultaneous Camera + Mic + GPS access in background", "CRITICAL (Score >= 100)"],
            ["TC-04", "Legitimate GMS Location", "Google Play Services location broker request", "BENIGN (Score = 0)"]
        ]
    )

    doc.save(doc_path)
    print("Updated docx successfully with complete Chapter 3!")

if __name__ == '__main__':
    update_chapter_3_properly()

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

def update_docx():
    doc_path = '/home/naantuam/Documents/Template for the Entire Project.docx'
    doc = docx.Document(doc_path)

    # Locate Chapter 3 start (P312) and Chapter 4 start (P328)
    p312_idx = None
    p328_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().upper() == 'CHAPTER THREE':
            p312_idx = i
        elif p.text.strip().upper() == 'CHAPTER FOUR' and p312_idx is not None:
            p328_idx = i
            break

    print(f"Chapter 3 start index: {p312_idx}, Chapter 4 start index: {p328_idx}")

    # Read paragraphs before Chapter 3 and after Chapter 4 placeholder
    # We will construct a new document or modify paragraphs in place.
    # To keep exact formatting and avoid breaking docx structure, let's delete the old placeholder paragraphs (P312 to P327) and insert new paragraphs before P328.

    # Reference paragraph to insert before
    target_p = doc.paragraphs[p328_idx]

    def add_p(text, style='Body Text', align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, space_after=6):
        p = target_p.insert_paragraph_before(text)
        p.alignment = align
        if style in doc.styles:
            p.style = doc.styles[style]
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold and len(p.runs) > 0:
            for r in p.runs:
                r.bold = True
        return p

    def add_heading_1(text):
        p = target_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_heading_2(text):
        p = target_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_heading_3(text):
        p = target_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_body(text):
        p = target_p.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 0, 0)
        return p

    # Insert Chapter 3 Content
    add_heading_1("CHAPTER THREE")
    add_heading_2("SYSTEM ANALYSIS AND DESIGN (MATERIALS AND METHODS)")

    add_heading_3("3.1 Introduction")
    add_body("This chapter presents the system analysis, architectural design, materials, and operational methods used to build and evaluate the hybrid sensor monitoring framework. Smartphones running the Android operating system rely on hardware peripherals such as camera lenses, microphones, location services, and motion sensors to deliver interactive application features. While modern mobile operating systems require user permission to access camera and microphone hardware, background data collection, side-channel sensor polling, and covert data transmission remain challenging security concerns. Existing detection tools often require superuser root privileges, perform heavy on-device processing that drains device battery, or depend completely on continuous internet connectivity. To address these limitations, this research introduces an unobtrusive hybrid monitoring framework that decouples lightweight mobile sensor auditing from real-time threat evaluation on an edge analysis server.")
    add_body("The primary objective of this chapter is to explain how the system was analyzed, designed, implemented, and evaluated. Specifically, the system development and empirical telemetry benchmarking were carried out over a twelve-week period from May 2026 to July 2026. All experimental data collection and testing procedures were conducted within the Cybersecurity Science Laboratory at the Federal University of Technology Minna, utilizing a dedicated hardware testbed connected via a secure local wireless network and USB tethering environment.")

    add_heading_3("3.2 Research Methodology Framework")
    add_body("The development of the hybrid sensor monitoring framework follows the Design Science Research Methodology (DSRM), a structured engineering framework designed for creating and evaluating cybersecurity software artifacts. The research methodology is organized into six continuous, sequential stages: problem identification and requirement gathering, system architecture topology design, mobile telemetry agent engineering, contextual threat rules matrix calibration, dual-database persistence design, and forensic web dashboard development.")
    add_body("In the first stage, empirical limitations of current Android sensor security tools were documented, establishing the core functional requirements for zero-root permission execution, real-time threat detection, and low battery consumption. In the second stage, a decoupled system topology was designed to separate data collection on the smartphone from heavy analytics on an edge server. In the third stage, an event-driven mobile auditing agent was created to query system services across multiple execution modes. In the fourth stage, a four-technique contextual rules engine was implemented and mapped to the MITRE ATT&CK Mobile framework. In the fifth stage, a dual-database architecture combining local SQLite buffering with a central PostgreSQL database was engineered to guarantee offline resilience. In the sixth stage, an interactive forensic dashboard was developed to visualize real-time risk scores and threat logs.")
    add_body("This structured methodology directly supports the fourth objective of this project: to evaluate the framework by mapping observed sensor behaviors to relevant techniques in the MITRE ATT&CK Mobile Matrix for structured threat classification using a custom rules engine, simulated attack testing, and live sensor monitoring across camera, microphone, location, and motion sensors.")

    add_heading_3("3.3 System Materials and Experimental Setup")
    add_body("The materials and tools used in this study comprise hardware components, software development packages, and live mobile sensor telemetry streams. The hardware testbed consists of a physical target smartphone (Infinix X683 running Android 10, API Level 29) and an edge host workstation powered by an Intel Core i7 processor running Linux. Communication between the mobile handset and host workstation is maintained through high-speed Wi-Fi and USB tethered interfaces.")
    add_body("The software stack includes Node.js as the runtime engine for the edge analysis server, React and Vite for the web dashboard interface, SQLite for local on-device buffering, and Neon PostgreSQL for centralized storage. Mobile telemetry is gathered directly from core Android system services, including sensorservice for accelerometer and light data, audio for microphone state, camera for camera service sessions, and location for GPS and network location queries.")

    add_heading_3("3.4 System Analysis and Requirement Specifications")
    add_body("System analysis was conducted to establish clear functional and non-functional requirements for the hybrid monitoring architecture. Functional requirements mandate that the framework must audit camera, microphone, location, and motion sensors without requiring root superuser permissions; evaluate telemetry events in real time against a multi-technique threat matrix; throttle repetitive alert logs to prevent user fatigue; and maintain full operation during network disconnections through local data buffering.")
    add_body("Non-functional requirements emphasize system efficiency and reliability. The mobile monitoring agent must consume less than five percent of total device battery capacity during continuous operation. The communication pipeline between the handset and the edge analysis server must maintain an event transmission latency below one hundred milliseconds. Furthermore, the web dashboard must render live telemetry events dynamically without causing user interface responsiveness delays.")

    add_heading_3("3.5 System Architecture and Operational Procedures")
    add_body("The hybrid monitoring architecture is built around a decoupled client-server model that separates lightweight data gathering from computationally intensive threat analysis. The system consists of three main operational tiers: the Mobile Telemetry Agent, the Edge Analysis Server, and the Interactive Forensic Dashboard.")
    add_body("The Mobile Telemetry Agent runs on or alongside the Android device and periodically queries system diagnostic interfaces. To ensure flexibility across different deployment scenarios, the agent supports two execution modes: Host-Side Wireless ADB Bridge mode and Native On-Device Termux mode. In Wireless ADB Bridge mode, the host workstation executes Android Debug Bridge commands over a local network to query device state without installing additional software on the phone. In Native Termux mode, the agent runs directly within the mobile terminal shell.")
    add_body("When the agent collects telemetry events, it packages them into structured JSON payloads containing package identifiers, sensor names, application states (foreground or background), display status (screen on or screen off), and active hardware flags. These events are transmitted asynchronously over a persistent WebSocket connection to the Edge Analysis Server for scoring and storage.")

    add_heading_3("3.6 Technical Threat Engine & MITRE ATT&CK Mapping")
    add_body("The core of the edge analysis server is a four-technique threat evaluation engine designed to assess the risk of every incoming sensor event. Rather than relying on simple binary flags, the engine computes a composite risk score based on application trust levels, functional coherence, multi-sensor correlation, and origin risk.")
    add_body("Technique 1 (App Trust Tiers) categorizes applications into trust levels that modify the base risk score. Operating system infrastructure packages (such as Google Play Services) receive a trust reduction of minus fifty points because background sensor access is architecturally expected. Trusted system apps receive a reduction of minus thirty points, while popular known apps receive minus fifteen points.")
    add_body("Technique 2 (Sensor-App Coherence) evaluates whether an application's sensor access aligns with its declared category. For example, a navigation app accessing GPS location receives a minus ten coherence match reward. Conversely, an application accessing an unexpected sensor (such as a financial app accessing the microphone) receives a plus fifteen mismatch penalty.")
    add_body("Technique 3 (Multi-Sensor Correlation) detects compound surveillance patterns. When an application accesses both camera and microphone simultaneously, an Audio-Video sync penalty of plus thirty-five points is applied. If camera, microphone, and location are accessed together, a full surveillance triad penalty of plus fifty points is triggered.")
    add_body("Technique 4 (Origin Risk) checks the installation source of unknown applications, applying a plus twenty points penalty to sideloaded packages installed from third-party APK files.")
    add_body("To prevent system location brokers like Google Play Services (GMS) from generating false alerts during routine location updates, the engine enforces a five-property OS Infrastructure Exemption. System brokers are exempted from background violation penalties because they possess platform signing certificates, act strictly as API brokers for other apps, undergo Play Protect verification, are tracked by system app-ops, and lack exfiltration data paths.")
    add_body("The composite risk score determines the final threat level: Critical (score of 100 or higher), High (score from 60 to 99), Suspicious (score from 25 to 59), and Benign (score below 25). Table 3.1 details how these rules map to the MITRE ATT&CK Mobile framework.")

    add_heading_3("Table 3.1: Hybrid Sensor Monitoring Framework Threat Matrix Mapped to MITRE ATT&CK Mobile")

    # Add Table 3.1
    table_data = [
        ["Tactical Domain", "Rule ID", "Sensor Behavior & Trigger Condition", "MITRE ID", "MITRE Technique Name", "Points"],
        ["Collection", "COLLECTION_MIC", "Unsanctioned or active microphone recording stream", "T1430", "Audio Capture", "+20"],
        ["Collection", "COLLECTION_CAMERA", "Unsanctioned or active camera session", "T1125", "Video Capture", "+20"],
        ["Collection", "COLLECTION_LOW_FREQ", "Motion/ambient light sensor polling (<= 20 Hz)", "T1636", "Sensor Information Discovery", "+5"],
        ["Collection", "COLLECTION_HIGH_FREQ", "Ultra-high frequency sensor polling (> 100 Hz)", "T1429", "Side-Channel Exploitation", "+20"],
        ["Discovery", "DISCOVERY_GPS", "GPS / Fine Location provider polling", "T1636", "Location Tracking", "+15"],
        ["Discovery", "DISCOVERY_BLE", "Bluetooth Low Energy (BLE) background scanning", "T1636", "Sensor Information Discovery", "+15"],
        ["Discovery", "DISCOVERY_FUSION", "Simultaneous motion sensor reads (Gyro/Accel)", "T1427", "System Information Discovery", "+20"],
        ["Context", "STATE_BACKGROUND", "Sensor access while app is in background", "N/A", "Contextual State Penalty", "+25"],
        ["Context", "STATE_SCREEN_OFF", "Sensor telemetry streaming while display is OFF", "N/A", "Covert Execution Context", "+30"],
        ["Context", "CONTEXT_BG_VIOLATION", "High-value sensor accessed in background", "N/A", "Unsanctioned Access", "+50"],
        ["Defense Evasion", "EVASION_STATE_DISCREPANCY", "Foreground claim mismatches display OFF state", "T1036", "Masquerading / State Mismatch", "+20"],
        ["Defense Evasion", "EVASION_ACCESSIBILITY", "Non-system package using Accessibility Services", "T1406", "Abuse Accessibility Features", "+15"],
        ["Defense Evasion", "EVASION_LOG_DELETE", "Local audit trail or log file deletion attempted", "T1403", "Indicator Removal on Host", "+30"],
        ["Exfiltration", "EXFIL_STAGING", "Large local JSON buffer staging detected", "T1430", "Data Staged", "+15"],
        ["Exfiltration", "EXFIL_IMMEDIATE", "Immediate data exfiltration over C2 channel", "T1041", "Exfiltration Over C2 Channel", "+40"],
        ["Correlation", "CORRELATION_AV_SYNC", "Camera and Microphone active simultaneously", "T1512", "Audio & Video Capture Sync", "+35"],
        ["Correlation", "CORRELATION_TRACK_TRIO", "Camera + Mic + Location active simultaneously", "T1430", "Full Surveillance Profile", "+50"],
        ["Correlation", "CORRELATION_NET_SENSOR", "Sensor collection concurrent with active exfiltration", "T1041", "Exfiltrated Sensor Capture", "+25"],
        ["Origin Risk", "ORIGIN_SIDELOADED", "Package installed from non-Play Store origin", "T1476", "Sideloaded Application", "+20"]
    ]

    t = doc.add_table(rows=len(table_data), cols=6)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Move table before target_p
    target_p._p.addprevious(t._tbl)

    for r_idx, row in enumerate(t.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.text = table_data[r_idx][c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [1, 3, 5] else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10 if r_idx > 0 else 10)
                if r_idx == 0:
                    run.bold = True

    add_heading_3("3.7 Database Design and Synchronization Protocols")
    add_body("To ensure high reliability and zero data loss during network outages, the framework uses a dual-database storage design combining local on-device storage with a central edge database.")
    add_body("On the mobile side, a lightweight SQLite database acts as a local buffer. If network connectivity to the edge server is lost, telemetry events and threat logs are saved locally in SQLite tables. Once connection is restored, an automated background synchronization process uploads the buffered events to the central Neon PostgreSQL database and clears the local queue. On the edge server side, PostgreSQL maintains long-term session records, active device profiles, and historical threat logs, enabling dynamic visualization on the forensic web dashboard.")

    # Now remove the old placeholder paragraphs (P312 to P327)
    # Re-query paragraphs in document to get current indices
    paragraphs_to_remove = []
    found_ch3 = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == 'CHAPTER THREE':
            found_ch3 = True
            paragraphs_to_remove.append(p)
        elif found_ch3:
            if txt == 'CHAPTER FOUR':
                break
            # Check if this is one of our newly added paragraphs or an old one
            # Old placeholder texts: 'MATERIALS AND METHODS OR SYSTEM ANALYSIS AND DESIGN', 'This chapter is the same...', 'WHEN the study...', etc.
            if any(k in txt for k in ['MATERIALS AND METHODS OR SYSTEM ANALYSIS AND DESIGN', 'This chapter is the same as research methodology', 'WHEN the study was carried out', 'WHERE the study was carried out', 'WHAT materials', 'HOW the study', 'WHAT procedures', 'NOTE!!!!!!!!', 'If your project is Research Based']):
                paragraphs_to_remove.append(p)

    print(f"Removing {len(paragraphs_to_remove)} placeholder paragraphs...")
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    doc.save(doc_path)
    print("Updated Word Document successfully!")

if __name__ == '__main__':
    update_docx()
